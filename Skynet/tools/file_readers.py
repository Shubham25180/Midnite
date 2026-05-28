"""File reading tools: PDF (with vision OCR fallback), spreadsheets, and Word documents."""
from __future__ import annotations

import base64
import csv
import json
import logging
import urllib.request
from pathlib import Path

logger = logging.getLogger(__name__)


def _vision_model() -> str:
    try:
        import yaml
        with open("config/settings.yaml", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        return cfg.get("vision", {}).get("model", "qwen2.5vl:7b")
    except Exception:
        return "qwen2.5vl:7b"


def read_pdf(path: str, pages: str = "all", dpi: int = 200) -> str:
    try:
        import fitz
    except ImportError:
        return "PDF reading requires pymupdf: pip install pymupdf"

    p = Path(path)
    if not p.exists():
        return f"File not found: {path}"
    if p.suffix.lower() != ".pdf":
        return f"Not a PDF file: {path}"

    # Clamp DPI: 150 minimum (below this OCR quality degrades badly),
    # 400 maximum (diminishing returns + very slow on large documents).
    dpi = max(150, min(400, dpi))

    try:
        doc = fitz.open(str(p))
    except Exception as exc:
        return f"Could not open PDF: {exc}"

    total = len(doc)
    indices = _parse_page_range(pages, total)

    page_texts: dict[int, str] = {}
    needs_ocr: list[int] = []

    for i in indices:
        text = doc[i].get_text().strip()
        if len(text) >= 80:
            page_texts[i] = text
        else:
            needs_ocr.append(i)

    if needs_ocr:
        logger.info("PDF %s: %d page(s) need vision OCR at %d DPI: %s", p.name, len(needs_ocr), dpi, needs_ocr)
        ocr = _ocr_pages(doc, needs_ocr, dpi)
        page_texts.update(ocr)

    doc.close()

    parts = [f"--- Page {i + 1} ---\n{page_texts.get(i, '(empty)')}" for i in indices]
    body = "\n\n".join(parts)

    if needs_ocr and len(needs_ocr) < len(indices):
        method = f"mixed (text + vision OCR @ {dpi} DPI)"
    elif needs_ocr:
        method = f"vision OCR @ {dpi} DPI"
    else:
        method = "text extraction"

    header = f"[PDF: {p.name}  pages={pages}  {total} total pages  method={method}]\n\n"
    result = header + body
    if len(result) > 8000:
        result = result[:8000] + "\n... (truncated)"
    return result


def _parse_page_range(pages: str, total: int) -> list[int]:
    if pages == "all":
        return list(range(total))
    if "-" in pages:
        start_str, end_str = pages.split("-", 1)
        start = max(0, int(start_str.strip()) - 1)
        end = min(total, int(end_str.strip()))
        return list(range(start, end))
    try:
        n = int(pages.strip())
        idx = n - 1
        if 0 <= idx < total:
            return [idx]
    except ValueError:
        pass
    return list(range(total))


def _ocr_pages(doc, indices: list[int], dpi: int = 200) -> dict[int, str]:
    """Render PDF pages as PNG and send to vision model for OCR."""
    model = _vision_model()
    results: dict[int, str] = {}

    for i in indices:
        try:
            pixmap = doc[i].get_pixmap(dpi=dpi)
            img_b64 = base64.b64encode(pixmap.tobytes("png")).decode()

            payload = {
                "model": model,
                "messages": [{
                    "role": "user",
                    "content": (
                        "Extract all text from this document page exactly as written. "
                        "Return only the text content, no commentary."
                    ),
                    "images": [img_b64],
                }],
                "stream": False,
            }
            data = json.dumps(payload).encode()
            req = urllib.request.Request(
                "http://localhost:11434/api/chat",
                data=data,
                headers={"Content-Type": "application/json"},
            )
            with urllib.request.urlopen(req, timeout=90) as resp:
                response = json.loads(resp.read())

            ocr_text = response.get("message", {}).get("content", "(no output)")
            results[i] = ocr_text
            logger.info("OCR page %d via %s: %d chars", i + 1, model, len(ocr_text))

        except Exception as exc:
            logger.warning("OCR failed page %d: %s", i + 1, exc)
            results[i] = f"(OCR failed: {exc})"

    return results


# ── Spreadsheet / CSV ──────────────────────────────────────────────────────────

def read_spreadsheet(path: str, sheet: str | None = None, max_rows: int = 100) -> str:
    p = Path(path)
    if not p.exists():
        return f"File not found: {path}"

    suffix = p.suffix.lower()
    if suffix == ".csv":
        return _read_csv(p, max_rows)
    if suffix in (".xlsx", ".xls", ".xlsm"):
        return _read_excel(p, sheet, max_rows)
    return f"Unsupported format: '{suffix}'. Supported: .csv, .xlsx, .xls, .xlsm"


def _read_csv(p: Path, max_rows: int) -> str:
    try:
        with open(p, encoding="utf-8", errors="replace", newline="") as f:
            rows = list(csv.reader(f))
    except Exception as exc:
        return f"CSV read error: {exc}"

    if not rows:
        return "Empty CSV file."

    headers = rows[0]
    data = rows[1:]
    total = len(data)
    preview = data[:max_rows]

    lines = [
        f"[CSV: {p.name}  {len(headers)} columns  {total} data rows]",
        "Columns: " + ", ".join(headers),
        f"\nFirst {min(max_rows, total)} rows:",
    ]
    for row in preview:
        lines.append("  " + " | ".join(str(v)[:50] for v in row))
    if total > max_rows:
        lines.append(f"\n... ({total - max_rows} more rows)")
    return "\n".join(lines)


def _read_excel(p: Path, sheet: str | None, max_rows: int) -> str:
    try:
        import openpyxl
    except ImportError:
        return "Excel reading requires openpyxl: pip install openpyxl"

    try:
        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
    except Exception as exc:
        return f"Could not open Excel file: {exc}"

    sheet_names = wb.sheetnames
    if sheet:
        if sheet not in sheet_names:
            wb.close()
            return f"Sheet '{sheet}' not found. Available: {', '.join(sheet_names)}"
        ws = wb[sheet]
    else:
        ws = wb.active
    selected = ws.title

    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        return f"Sheet '{selected}' is empty."

    headers = [str(v) if v is not None else "" for v in rows[0]]
    data = rows[1:]
    total = len(data)
    preview = data[:max_rows]

    lines = [
        f"[Excel: {p.name}  sheet='{selected}'  {len(headers)} columns  {total} data rows]",
        f"Available sheets: {', '.join(sheet_names)}",
        "Columns: " + ", ".join(headers),
        f"\nFirst {min(max_rows, total)} rows:",
    ]
    for row in preview:
        lines.append("  " + " | ".join((str(v)[:40] if v is not None else "") for v in row))
    if total > max_rows:
        lines.append(f"\n... ({total - max_rows} more rows)")
    return "\n".join(lines)


# ── Word documents ─────────────────────────────────────────────────────────────

def read_word(path: str) -> str:
    """Read a .docx Word document and return formatted text content."""
    try:
        from docx import Document
    except ImportError:
        return "Word reading requires python-docx: pip install python-docx"

    p = Path(path)
    if not p.exists():
        return f"File not found: {path}"
    if p.suffix.lower() not in (".docx",):
        return f"Not a .docx file: {path}  (legacy .doc format is not supported)"

    try:
        doc = Document(str(p))
    except Exception as exc:
        return f"Could not open Word document: {exc}"

    lines: list[str] = []

    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        style = para.style.name
        if style.startswith("Heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
            lines.append("#" * level + " " + para.text)
        elif style.startswith("List"):
            lines.append("- " + para.text)
        else:
            lines.append(para.text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n[Table {i + 1}]")
        for row in table.rows:
            lines.append("  " + " | ".join(cell.text.strip()[:40] for cell in row.cells))

    content = "\n".join(lines)
    if len(content) > 8000:
        content = content[:8000] + "\n... (truncated)"

    return f"[Word: {p.name}  {len(doc.paragraphs)} paragraphs  {len(doc.tables)} tables]\n\n{content}"
