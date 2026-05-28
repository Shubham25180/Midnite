"""File writing tools: spreadsheets (CSV/XLSX) and Word documents (.docx)."""
from __future__ import annotations

import csv
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ── Spreadsheet writing ────────────────────────────────────────────────────────

def write_spreadsheet(path: str, rows: list, sheet: str = "Sheet1") -> str:
    """
    Write structured data to a CSV or XLSX file.

    rows: list of dicts — keys from the first dict become column headers.
    Example: [{"Name": "Alice", "Score": 95}, {"Name": "Bob", "Score": 88}]

    For .xlsx: bold headers, auto column widths.
    For .csv: standard comma-separated, UTF-8.
    """
    if not rows:
        return "No data provided — rows list is empty."

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    suffix = p.suffix.lower()

    if not isinstance(rows[0], dict):
        return (
            "rows must be a list of dicts, e.g. "
            '[{"Name": "Alice", "Score": 95}, {"Name": "Bob", "Score": 88}]'
        )

    headers = list(rows[0].keys())

    if suffix == ".csv":
        return _write_csv(p, rows, headers)
    if suffix in (".xlsx", ".xlsm"):
        return _write_xlsx(p, rows, headers, sheet)
    return f"Unsupported format: '{suffix}'. Use .csv or .xlsx"


def _write_csv(p: Path, rows: list, headers: list) -> str:
    try:
        with open(p, "w", encoding="utf-8", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=headers, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(rows)
        logger.info("write_spreadsheet CSV: %s  (%d rows)", p, len(rows))
        return f"Written: {p}  ({len(rows)} rows, {len(headers)} columns)"
    except Exception as exc:
        return f"CSV write error: {exc}"


def _write_xlsx(p: Path, rows: list, headers: list, sheet: str) -> str:
    try:
        import openpyxl
        from openpyxl.styles import Font
    except ImportError:
        return "XLSX writing requires openpyxl: pip install openpyxl"

    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet

        ws.append(headers)
        for cell in ws[1]:
            cell.font = Font(bold=True)

        for row in rows:
            ws.append([row.get(h, "") for h in headers])

        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        wb.save(str(p))
        logger.info("write_spreadsheet XLSX: %s  (%d rows, sheet=%r)", p, len(rows), sheet)
        return f"Written: {p}  ({len(rows)} rows, {len(headers)} columns, sheet='{sheet}')"
    except Exception as exc:
        return f"XLSX write error: {exc}"


# ── Word document writing ──────────────────────────────────────────────────────

def write_word(path: str, content: str) -> str:
    """
    Create a .docx Word document from structured plain text.

    Formatting syntax (one element per line):
      # Title          → Heading 1
      ## Section       → Heading 2
      ### Subsection   → Heading 3
      - item           → Bullet point
      1. item          → Numbered list
      | col | col |    → Table row  (use | --- | --- | as separator row)
      plain text       → Paragraph

    Example:
      # Sales Report
      ## Q1 Summary
      Revenue was up 12% quarter-on-quarter.
      - Strong performance in Europe
      - New product line contributed 30%
      | Region | Revenue |
      | --- | --- |
      | Europe | $1.2M |
      | Asia | $0.9M |
    """
    try:
        from docx import Document
    except ImportError:
        return "Word creation requires python-docx: pip install python-docx"

    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)
    if p.suffix.lower() != ".docx":
        return f"Output must be a .docx file, got: {path}"

    try:
        doc = Document()
        _parse_content_into_doc(doc, content)
        doc.save(str(p))
        para_count = sum(1 for para in doc.paragraphs if para.text.strip())
        logger.info("write_word: %s  (%d paragraphs)", p, para_count)
        return f"Written: {p}  ({para_count} paragraphs)"
    except Exception as exc:
        return f"Word write error: {exc}"


def _parse_content_into_doc(doc, content: str) -> None:
    """Parse line-by-line content and add elements to a python-docx Document."""
    lines = content.splitlines()
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        if not table_buffer:
            return
        # Drop separator rows (cells that are all dashes)
        data = [r for r in table_buffer if not all(c.strip("-") == "" for c in r)]
        if not data:
            table_buffer.clear()
            return
        ncols = max(len(r) for r in data)
        t = doc.add_table(rows=len(data), cols=ncols)
        t.style = "Table Grid"
        for ri, row in enumerate(data):
            for ci in range(min(len(row), ncols)):
                t.cell(ri, ci).text = row[ci].strip()
        table_buffer.clear()

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped[1:-1].split("|")]
            table_buffer.append(cells)
            continue
        else:
            flush_table()

        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in ".)" and stripped[2] == " ":
            doc.add_paragraph(stripped[3:], style="List Number")
        else:
            doc.add_paragraph(stripped)

    flush_table()
